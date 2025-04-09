from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import sys

def clean_text(text):
    """清理非法字符和多余空格"""
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def set_page_settings(doc):
    """设置页面布局"""
    for section in doc.sections:
        # A4纸张大小 (21.0×29.7cm)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        
        # 页边距（1厘米）
        section.top_margin = Cm(1)
        section.right_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        
        # 页眉/页脚距离
        section.header_distance = Cm(0.1)
        section.footer_distance = Cm(0.4)

def set_columns(section, num_columns=2, space_between=0.5):
    """设置分栏"""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), str(int(space_between * 567)))

def add_footer(doc):
    """添加页脚（共X页 第Y页）"""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run("共")
        
        # 总页数
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText1 = OxmlElement('w:instrText')
        instrText1.text = 'NUMPAGES'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText1, fldChar2])
        
        run.add_text("页 第")
        
        # 当前页码
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'PAGE'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar3, instrText2, fldChar4])
        
        run.add_text("页")
        
        # 设置页脚字体
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'  # 英文名称
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 中文名称
            run.font.size = Pt(8)

def set_font(run):
    """确保中英文字体都设置为微软雅黑"""
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(8)

def txt_to_docx(input_txt, output_docx):
    try:
        with open(input_txt, 'rb') as f:
            text = f.read().decode('utf-8', errors='ignore')
        
        text = clean_text(text)
        text = text.replace('\n', '').replace('\r', '')
        
        doc = Document()
        set_page_settings(doc)
        
        for section in doc.sections:
            set_columns(section)
        
        add_footer(doc)
        
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_format.line_spacing = Pt(9)
        
        for run in p.runs:
            set_font(run)
        
        doc.save(output_docx)
        print(f"转换成功: {input_txt} -> {output_docx}")
        print("提示：如果字体仍未生效，请尝试在Word中手动刷新字体缓存（关闭后重新打开文档）")
    
    except Exception as e:
        print(f" 转换失败: {e}")
        if "font" in str(e).lower():
            print("字体问题：请确认系统已安装微软雅黑字体")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        input_file = input("请输入TXT文件路径: ").strip('"')
    
    output_file = input_file.rsplit('.', 1)[0] + '.docx'
    txt_to_docx(input_file, output_file)
    # 等待用户按键退出
    input("操作完成，按任意键退出...")
